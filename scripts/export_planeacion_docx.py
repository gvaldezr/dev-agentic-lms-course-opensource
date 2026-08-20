from __future__ import annotations

import argparse
import json
import re
from datetime import date, timedelta
from pathlib import Path
from typing import Any

from docx import Document


def _num(value: Any, default: Any = "") -> Any:
    if value is None or value == "":
        return default
    return value


def _add_list(document: Document, items: list[str]) -> None:
    for item in items:
        text = str(item).strip()
        if text:
            try:
                document.add_paragraph(text, style="List Bullet")
            except KeyError:
                document.add_paragraph(f"- {text}")


def _phase_items(session: dict[str, Any], phase: str) -> list[str]:
    key = f"{phase}_actividades"
    if isinstance(session.get(key), list) and session[key]:
        return [str(x) for x in session[key] if str(x).strip()]

    block = session.get(phase)
    if isinstance(block, dict):
        instructions = str(block.get("instrucciones", "")).strip()
        if instructions:
            return [instructions]
    return []


def _minutes_per_session(reglas: dict[str, Any]) -> int:
    raw = reglas.get("minutos_por_sesion")
    if isinstance(raw, list) and raw:
        nums = [int(x) for x in raw if str(x).strip().isdigit()]
        if nums:
            return nums[0]
    if isinstance(raw, (int, float)):
        return int(raw)
    if isinstance(raw, str) and raw.isdigit():
        return int(raw)
    return 90


def _join_or_dash(items: list[str]) -> str:
    cleaned = [str(i).strip() for i in items if str(i).strip()]
    return "\n".join(cleaned) if cleaned else "-"


def _safe_set(table: Any, row: int, col: int, value: Any) -> None:
    if row < 0 or col < 0:
        return
    if row >= len(table.rows):
        return
    if col >= len(table.rows[row].cells):
        return
    table.rows[row].cells[col].text = str(value)


def _flatten_adas(periodos: list[dict[str, Any]]) -> list[dict[str, Any]]:
    adas: list[dict[str, Any]] = []
    for periodo in periodos:
        adas.extend(periodo.get("adas") or [])
    return adas


def _weeks_span(ada: dict[str, Any]) -> str:
    semanas = sorted(
        {
            s.get("semana")
            for s in (ada.get("sesiones_desarrolladas") or [])
            if s.get("semana") is not None
        }
    )
    if not semanas:
        return "S/N"
    if len(semanas) == 1:
        return f"Semana {semanas[0]}"
    return f"Semanas {semanas[0]}-{semanas[-1]}"


def _week_bounds(ada: dict[str, Any]) -> tuple[int | None, int | None]:
    semanas = sorted(
        {
            s.get("semana")
            for s in (ada.get("sesiones_desarrolladas") or [])
            if s.get("semana") is not None
        }
    )
    if not semanas:
        return None, None
    return int(semanas[0]), int(semanas[-1])


def _parse_academic_date(value: str | None) -> date | None:
    if not value:
        return None
    raw = value.strip()
    if not raw:
        return None

    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return date.fromisoformat(raw) if fmt == "%Y-%m-%d" else date.strptime(raw, fmt)  # type: ignore[attr-defined]
        except Exception:
            pass

    # Fallback seguro sin date.strptime para compatibilidad amplia
    m = re.match(r"^(\d{1,2})[/-](\d{1,2})[/-](\d{4})$", raw)
    if m:
        d, mo, y = map(int, m.groups())
        return date(y, mo, d)

    # Formato: 17 de agosto de 2026
    m2 = re.match(r"^(\d{1,2})\s+de\s+([a-zA-ZáéíóúÁÉÍÓÚ]+)\s+de\s+(\d{4})$", raw)
    if not m2:
        return None

    day = int(m2.group(1))
    month_name = m2.group(2).lower()
    month_name = (
        month_name.replace("á", "a")
        .replace("é", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ú", "u")
    )
    year = int(m2.group(3))

    month_map = {
        "enero": 1,
        "febrero": 2,
        "marzo": 3,
        "abril": 4,
        "mayo": 5,
        "junio": 6,
        "julio": 7,
        "agosto": 8,
        "septiembre": 9,
        "setiembre": 9,
        "octubre": 10,
        "noviembre": 11,
        "diciembre": 12,
    }
    month = month_map.get(month_name)
    if not month:
        return None
    return date(year, month, day)


def _week_to_date(start: date, week: int) -> date:
    return start + timedelta(days=(max(week, 1) - 1) * 7)


def _format_date(value: date) -> str:
    return value.strftime("%d/%m/%Y")


def _unit_dates_text(
    ada_1: dict[str, Any],
    ada_2: dict[str, Any],
    start_date: date | None,
    end_date: date | None,
) -> tuple[str, str]:
    if not start_date:
        return _weeks_span(ada_1) if ada_1 else "S/N", _weeks_span(ada_2) if ada_2 else "S/N"

    w1s, w1e = _week_bounds(ada_1) if ada_1 else (None, None)
    w2s, w2e = _week_bounds(ada_2) if ada_2 else (None, None)

    # Fecha de entrega: cierre del ultimo ADA de la unidad
    last_week = w2e or w1e
    if last_week is None:
        return "S/N", "S/N"
    entrega = _week_to_date(start_date, last_week)
    if end_date and entrega > end_date:
        entrega = end_date

    # Retroalimentacion: una semana despues (o tope de fin de curso)
    retro = entrega + timedelta(days=7)
    if end_date and retro > end_date:
        retro = end_date

    return _format_date(entrega), _format_date(retro)


def _template_units_from_adas(adas_flat: list[dict[str, Any]]) -> list[dict[str, Any]]:
    units: list[dict[str, Any]] = []
    chunks = [adas_flat[0:2], adas_flat[2:4], adas_flat[4:6]]
    for idx, chunk in enumerate(chunks, start=1):
        non_empty = [a for a in chunk if a]
        objetivo = non_empty[0].get("objetivo", "") if non_empty else ""
        sessions = sum(int(_num(a.get("duracion_sesiones"), 0) or 0) for a in non_empty)
        units.append(
            {
                "periodo": idx,
                "nombre": f"Unidad {idx}",
                "objetivo": objetivo,
                "adas": non_empty,
                "duracion_sesiones": sessions,
            }
        )
    return units


def _resultados_text(ada: dict[str, Any]) -> str:
    resultados = ada.get("resultados_aprendizaje") or []
    if resultados:
        return _join_or_dash([str(x) for x in resultados])
    if ada.get("resultado_aprendizaje"):
        return str(ada.get("resultado_aprendizaje"))
    return "-"


def _fill_unit_table(table: Any, periodo: dict[str, Any], periodo_num: int, competencia: str, min_sesion: int) -> None:
    adas = periodo.get("adas") or []
    sesiones = sum(int(_num(a.get("duracion_sesiones"), 0) or 0) for a in adas)
    horas = round((sesiones * min_sesion) / 60.0, 1)

    for col in range(min(4, len(table.rows[0].cells))):
        _safe_set(table, 0, col, f"SECUENCIA DIDACTICA UNIDAD {periodo_num}")
    _safe_set(table, 1, 0, f"Unidad {periodo_num}")
    _safe_set(table, 1, 1, periodo.get("nombre", ""))
    _safe_set(table, 2, 0, "Competencia de la asignatura")
    _safe_set(table, 2, 1, competencia)
    _safe_set(table, 3, 0, f"Competencia de unidad {periodo_num}")
    _safe_set(table, 3, 1, periodo.get("objetivo", ""))
    _safe_set(table, 4, 0, "Competencia del perfil de egreso a los que contribuye la unidad")
    _safe_set(table, 4, 1, periodo.get("perfil_egreso", "No especificado"))
    _safe_set(table, 5, 0, "Duracion en sesiones")
    _safe_set(table, 5, 1, sesiones)
    _safe_set(table, 5, 2, "Duracion en horas")
    _safe_set(table, 5, 3, horas)


def _fill_activity_table(table: Any, ada: dict[str, Any] | None, fallback_num: int) -> None:
    if not ada:
        _safe_set(table, 0, 0, f"Actividad No. {fallback_num}")
        _safe_set(table, 0, 2, "Proceso ( )")
        _safe_set(table, 0, 3, "Producto ( )")
        _safe_set(table, 1, 2, "0")
        _safe_set(table, 2, 2, "0 puntos")
        _safe_set(table, 3, 1, "-")
        _safe_set(table, 4, 1, "-")
        _safe_set(table, 5, 1, "-")
        return

    numero = _num(ada.get("actividad_numero"), fallback_num)
    nombre = ada.get("nombre") or f"Actividad No. {numero}"
    tipo = str(ada.get("tipo_actividad", "proceso")).strip().lower()
    _safe_set(table, 0, 0, nombre)
    _safe_set(table, 0, 2, "Proceso (X)" if tipo == "proceso" else "Proceso ( )")
    _safe_set(table, 0, 3, "Producto (X)" if tipo == "producto" else "Producto ( )")
    _safe_set(table, 1, 0, nombre)
    _safe_set(table, 1, 2, _num(ada.get("duracion_sesiones"), 0))
    _safe_set(table, 2, 0, nombre)
    _safe_set(table, 2, 2, f"{_num(ada.get('puntaje'),0)} puntos")
    _safe_set(table, 3, 0, "Resultado(s) de aprendizaje")
    _safe_set(table, 3, 1, _resultados_text(ada))
    _safe_set(table, 4, 0, "Contenidos")
    _safe_set(table, 4, 1, _join_or_dash([str(x) for x in (ada.get("contenidos_a_desarrollar") or [])]))
    _safe_set(table, 5, 0, "Estrategias de ensenanza y aprendizaje")
    _safe_set(table, 5, 1, _join_or_dash([str(x) for x in (ada.get("estrategias_ensenanza_aprendizaje") or [])]))


def _fill_evidence_table(table: Any, ada: dict[str, Any] | None) -> None:
    if not ada:
        _safe_set(table, 0, 1, "-")
        _safe_set(table, 1, 1, "Lista de cotejo")
        return
    _safe_set(table, 0, 0, "Evidencia de aprendizaje")
    _safe_set(table, 0, 1, _join_or_dash([str(x) for x in (ada.get("evidencias_aprendizaje") or [])]))
    _safe_set(table, 1, 0, "Instrumento de evaluacion")
    _safe_set(table, 1, 1, str(ada.get("instrumento_evaluacion", "Lista de cotejo")))


def build_doc_from_template(
    input_json: Path,
    template_docx: Path,
    output_docx: Path,
    start_date_text: str | None = None,
    end_date_text: str | None = None,
) -> Path:
    with input_json.open("r", encoding="utf-8") as f:
        data = json.load(f)

    document = Document(str(template_docx))

    course_name = data.get("course_name") or "Planeacion didactica"
    programa = data.get("programa_asignatura") or {}
    datos = programa.get("datos_generales") or {}
    planeacion = data.get("planeacion_operativa") or {}
    reglas = planeacion.get("reglas") or {}
    competencia = str(data.get("competencia_general") or "")
    periodos = (data.get("estructura_curso_adas") or {}).get("periodos") or []
    min_sesion = _minutes_per_session(reglas)
    adas_flat = _flatten_adas(periodos)
    template_units = _template_units_from_adas(adas_flat)
    start_date = _parse_academic_date(start_date_text)
    end_date = _parse_academic_date(end_date_text)

    # Ajuste de portada
    for p in document.paragraphs:
        txt = p.text.strip()
        if "(Nombre de la asignatura)" in txt:
            p.text = course_name

    # 1) Datos generales
    if len(document.tables) >= 5:
        t3 = document.tables[2]
        _safe_set(t3, 1, 1, str(datos.get("numero_creditos", "")))
        _safe_set(t3, 1, 3, str(datos.get("tipo_creditos", "")))

        dist = datos.get("distribucion_horas") or {}
        t4 = document.tables[3]
        _safe_set(t4, 2, 1, str(datos.get("duracion_total_horas", "")))
        _safe_set(t4, 2, 2, str(dist.get("HCP", "")))
        _safe_set(t4, 2, 3, str(dist.get("HEI", "")))
        _safe_set(t4, 2, 4, str(dist.get("practicas_formativas_laboratorios", "")))

        t5 = document.tables[4]
        _safe_set(t5, 0, 1, str(datos.get("ubicacion", "")))
        _safe_set(t5, 1, 1, str(datos.get("requisitos_previos", "")))

    # 2) Intencionalidad y 3) Competencia
    if len(document.tables) >= 7:
        _safe_set(document.tables[5], 0, 0, "2. Intencionalidad de la asignatura\n" + str(programa.get("contexto_asignatura", "")))
        _safe_set(document.tables[6], 0, 0, "3. Competencia de la asignatura\n" + competencia)

    # 4) Distribucion de puntajes
    if len(document.tables) >= 9:
        t9 = document.tables[8]
        total_puntaje = 0
        for i in range(3):
            adas = (template_units[i].get("adas") if i < len(template_units) else []) or []
            a1 = adas[0] if len(adas) > 0 else {}
            a2 = adas[1] if len(adas) > 1 else {}
            row_idx = 2 + i
            _safe_set(t9, row_idx, 0, f"Unidad {i+1}")
            _safe_set(t9, row_idx, 1, f"{a1.get('nombre','')} ({_num(a1.get('duracion_sesiones'),0)} sesiones)")
            _safe_set(t9, row_idx, 2, f"{a2.get('nombre','')} ({_num(a2.get('duracion_sesiones'),0)} sesiones)")
            entrega_txt, retro_txt = _unit_dates_text(a1, a2, start_date, end_date)
            _safe_set(t9, row_idx, 3, entrega_txt)
            _safe_set(t9, row_idx, 4, retro_txt)
            evs = []
            if a1:
                evs.append((a1.get("evidencias_aprendizaje") or [""])[0])
            if a2:
                evs.append((a2.get("evidencias_aprendizaje") or [""])[0])
            _safe_set(t9, row_idx, 5, _join_or_dash([e for e in evs if e]))
            punt = int(_num(a1.get("puntaje"), 0) or 0) + int(_num(a2.get("puntaje"), 0) or 0)
            total_puntaje += punt
            _safe_set(t9, row_idx, 6, str(punt))

        _safe_set(t9, 5, 0, "Evaluacion de producto")
        _safe_set(t9, 5, 1, "Integrada al cierre del proceso (ADA final)")
        _safe_set(t9, 5, 5, "Compilado final de evidencias y lista de cotejo")
        _safe_set(t9, 6, 6, str(total_puntaje))

    # 5) Secuencia didactica por unidad y actividades
    while len(adas_flat) < 6:
        adas_flat.append(None)  # type: ignore[arg-type]

    unit_table_indices = [10, 17, 24]  # tablas 11, 18, 25
    activity_table_indices = [11, 14, 18, 21, 25, 28]  # tablas 12,15,19,22,26,29
    evidence_table_indices = [12, 15, 19, 22, 26, 29]  # tablas 13,16,20,23,27,30

    for idx, t_idx in enumerate(unit_table_indices, start=1):
        if t_idx < len(document.tables):
            periodo = template_units[idx - 1] if idx - 1 < len(template_units) else {"nombre": ""}
            _fill_unit_table(document.tables[t_idx], periodo, idx, competencia, min_sesion)

    for idx, t_idx in enumerate(activity_table_indices, start=1):
        if t_idx < len(document.tables):
            _fill_activity_table(document.tables[t_idx], adas_flat[idx - 1], idx)

    for idx, t_idx in enumerate(evidence_table_indices, start=1):
        if t_idx < len(document.tables):
            _fill_evidence_table(document.tables[t_idx], adas_flat[idx - 1])

    # 6) Evaluacion de producto
    if len(document.tables) >= 34:
        t33 = document.tables[32]
        cierre = adas_flat[4] if len(adas_flat) > 4 and adas_flat[4] else {}
        _safe_set(t33, 0, 1, str(_num(cierre.get("duracion_sesiones"), 1)))
        _safe_set(t33, 0, 3, str(round((int(_num(cierre.get("duracion_sesiones"), 1)) * min_sesion) / 60.0, 1)))
        _safe_set(t33, 1, 2, str(_num(cierre.get("puntaje"), 10)))
        _safe_set(t33, 2, 1, competencia)
        _safe_set(t33, 3, 1, _join_or_dash([str(x) for x in (cierre.get("estrategias_ensenanza_aprendizaje") or [])]))

        t34 = document.tables[33]
        _safe_set(t34, 0, 1, _join_or_dash([str(x) for x in (cierre.get("evidencias_aprendizaje") or [])]))
        _safe_set(t34, 1, 1, str(cierre.get("instrumento_evaluacion", "Lista de cotejo")))

    # Anexo de guion por sesiones
    document.add_page_break()
    document.add_heading("ANEXO: Guion de clase por sesiones", level=1)
    for periodo in periodos:
        for ada in (periodo.get("adas") or []):
            document.add_heading(str(ada.get("nombre", "Actividad")), level=2)
            for ses in (ada.get("sesiones_desarrolladas") or []):
                ses_num = _num(ses.get("sesion"), "?")
                document.add_paragraph(f"SESION {ses_num}")
                document.add_paragraph(f"Tema: {str(ses.get('tema', '')).strip()}")
                obj_ses = str(ses.get("objetivo_especifico") or ses.get("objetivo") or "").strip()
                document.add_paragraph(f"Objetivo especifico: {obj_ses}")
                document.add_paragraph("Resumen de los datos esenciales:")
                resumen = ses.get("resumen_datos_esenciales") or []
                if resumen:
                    _add_list(document, [str(x) for x in resumen])
                else:
                    document.add_paragraph("-")

                for phase in ("inicio", "desarrollo", "cierre"):
                    document.add_paragraph(phase.upper())
                    items = _phase_items(ses, phase)
                    if items:
                        _add_list(document, items)
                    else:
                        document.add_paragraph("Sin actividades registradas.", style="List Bullet")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))
    return output_docx


def build_doc(input_json: Path, output_docx: Path) -> Path:
    with input_json.open("r", encoding="utf-8") as f:
        data = json.load(f)

    document = Document()

    course_name = data.get("course_name") or "Planeacion didactica"
    programa = data.get("programa_asignatura") or {}
    datos = programa.get("datos_generales") or {}
    competencia = data.get("competencia_general") or ""
    planeacion = data.get("planeacion_operativa") or {}
    reglas = planeacion.get("reglas") or {}
    estructura = data.get("estructura_curso_adas") or {}
    periodos = estructura.get("periodos") or []
    min_sesion = _minutes_per_session(reglas)

    document.add_heading("UNIVERSIDAD AUTONOMA DE YUCATAN", level=0)
    document.add_paragraph("PLANEACION DIDACTICA")
    document.add_paragraph("CURSO REGULAR")
    document.add_paragraph(str(course_name))

    document.add_heading("1. Datos generales de identificacion", level=1)
    t_modal = document.add_table(rows=1, cols=3)
    t_modal.rows[0].cells[0].text = "Modalidad de la asignatura"
    t_modal.rows[0].cells[1].text = "Presencial (X)"
    t_modal.rows[0].cells[2].text = "No presencial ( )"

    t_cred = document.add_table(rows=2, cols=2)
    t_cred.rows[0].cells[0].text = "Numero de creditos"
    t_cred.rows[0].cells[1].text = str(datos.get("numero_creditos", ""))
    t_cred.rows[1].cells[0].text = "Tipo de creditos"
    t_cred.rows[1].cells[1].text = str(datos.get("tipo_creditos", ""))

    dist = datos.get("distribucion_horas") or {}
    t_horas = document.add_table(rows=4, cols=2)
    t_horas.rows[0].cells[0].text = "Duracion total de la asignatura"
    t_horas.rows[0].cells[1].text = str(datos.get("duracion_total_horas", ""))
    t_horas.rows[1].cells[0].text = "HCP"
    t_horas.rows[1].cells[1].text = str(dist.get("HCP", ""))
    t_horas.rows[2].cells[0].text = "HEI"
    t_horas.rows[2].cells[1].text = str(dist.get("HEI", ""))
    t_horas.rows[3].cells[0].text = "Practicas formativas (HPF)"
    t_horas.rows[3].cells[1].text = str(dist.get("practicas_formativas_laboratorios", ""))

    t_ubic = document.add_table(rows=2, cols=2)
    t_ubic.rows[0].cells[0].text = "Ubicacion"
    t_ubic.rows[0].cells[1].text = str(datos.get("ubicacion", ""))
    t_ubic.rows[1].cells[0].text = "Requisitos academicos previos"
    t_ubic.rows[1].cells[1].text = str(datos.get("requisitos_previos", ""))

    document.add_heading("2. Intencionalidad de la asignatura", level=1)
    document.add_paragraph(str(programa.get("contexto_asignatura", "")))

    document.add_heading("3. Competencia de la asignatura", level=1)
    document.add_paragraph(str(competencia))

    document.add_heading("4. Acreditacion de la asignatura", level=1)
    t_dist = document.add_table(rows=1, cols=7)
    hdr = t_dist.rows[0].cells
    hdr[0].text = "Evaluacion de proceso"
    hdr[1].text = "ADA 1"
    hdr[2].text = "ADA 2"
    hdr[3].text = "Fecha de entrega"
    hdr[4].text = "Fecha de retroalimentacion"
    hdr[5].text = "Evidencia a entregar"
    hdr[6].text = "Puntaje"

    total_puntaje = 0
    for idx, periodo in enumerate(periodos, start=1):
        row = t_dist.add_row().cells
        row[0].text = f"Unidad {idx}"
        adas = periodo.get("adas") or []
        a1 = adas[0] if len(adas) > 0 else {}
        a2 = adas[1] if len(adas) > 1 else {}
        row[1].text = f"{a1.get('nombre','')} ({_num(a1.get('duracion_sesiones'),'0')} sesiones)"
        row[2].text = f"{a2.get('nombre','')} ({_num(a2.get('duracion_sesiones'),'0')} sesiones)"
        row[3].text = "Por definir"
        row[4].text = "Por definir"
        e1 = (a1.get("evidencias_aprendizaje") or [""])[0]
        row[5].text = str(e1)
        puntaje = int(_num(a1.get("puntaje"), 0) or 0) + int(_num(a2.get("puntaje"), 0) or 0)
        total_puntaje += puntaje
        row[6].text = str(puntaje)

    row_prod = t_dist.add_row().cells
    row_prod[0].text = "Evaluacion de producto"
    row_prod[1].text = "No aplica (configuracion actual)"
    row_prod[6].text = "0"

    row_total = t_dist.add_row().cells
    row_total[5].text = "Total"
    row_total[6].text = str(total_puntaje)

    document.add_heading("5. Secuencia didactica y Guion de clase", level=1)
    for p_idx, periodo in enumerate(periodos, start=1):
        periodo_num = _num(periodo.get("periodo"), p_idx)
        adas = periodo.get("adas") or []
        sesiones_periodo = sum(int(_num(a.get("duracion_sesiones"), 0) or 0) for a in adas)
        horas_periodo = round((sesiones_periodo * min_sesion) / 60.0, 1)

        t_unidad = document.add_table(rows=6, cols=4)
        for c in range(4):
            t_unidad.rows[0].cells[c].text = f"SECUENCIA DIDACTICA UNIDAD {periodo_num}"
        t_unidad.rows[1].cells[0].text = f"Unidad {periodo_num}"
        t_unidad.rows[1].cells[1].text = str(periodo.get("nombre", ""))
        t_unidad.rows[2].cells[0].text = "Competencia de la asignatura"
        t_unidad.rows[2].cells[1].text = str(competencia)
        t_unidad.rows[3].cells[0].text = f"Competencia de unidad {periodo_num}"
        t_unidad.rows[3].cells[1].text = str(periodo.get("objetivo", ""))
        t_unidad.rows[4].cells[0].text = "Competencia del perfil de egreso"
        t_unidad.rows[4].cells[1].text = str(periodo.get("perfil_egreso", "No especificado"))
        t_unidad.rows[5].cells[0].text = "Duracion en sesiones"
        t_unidad.rows[5].cells[1].text = str(sesiones_periodo)
        t_unidad.rows[5].cells[2].text = "Duracion en horas"
        t_unidad.rows[5].cells[3].text = str(horas_periodo)

        for a_idx, ada in enumerate(adas, start=1):
            numero = _num(ada.get("actividad_numero"), a_idx)
            nombre = ada.get("nombre") or f"Actividad No. {numero}"
            tipo = str(ada.get("tipo_actividad", "proceso")).strip().lower()

            t_act = document.add_table(rows=6, cols=4)
            t_act.rows[0].cells[0].text = nombre
            t_act.rows[0].cells[1].text = "Tipo"
            t_act.rows[0].cells[2].text = "Proceso (X)" if tipo == "proceso" else "Proceso ( )"
            t_act.rows[0].cells[3].text = "Producto (X)" if tipo == "producto" else "Producto ( )"
            t_act.rows[1].cells[0].text = nombre
            t_act.rows[1].cells[1].text = "Duracion en sesiones"
            t_act.rows[1].cells[2].text = str(_num(ada.get("duracion_sesiones"), ""))
            t_act.rows[2].cells[0].text = nombre
            t_act.rows[2].cells[1].text = "Puntaje"
            t_act.rows[2].cells[2].text = f"{_num(ada.get('puntaje'),'')} puntos"
            resultados = ada.get("resultados_aprendizaje") or []
            if not resultados and ada.get("resultado_aprendizaje"):
                resultados = [ada.get("resultado_aprendizaje")]
            t_act.rows[3].cells[0].text = "Resultado(s) de aprendizaje"
            t_act.rows[3].cells[1].text = _join_or_dash([str(x) for x in resultados])
            t_act.rows[4].cells[0].text = "Contenidos"
            t_act.rows[4].cells[1].text = _join_or_dash([str(x) for x in (ada.get("contenidos_a_desarrollar") or [])])
            t_act.rows[5].cells[0].text = "Estrategias de ensenanza y aprendizaje"
            t_act.rows[5].cells[1].text = _join_or_dash([str(x) for x in (ada.get("estrategias_ensenanza_aprendizaje") or [])])

            t_ev = document.add_table(rows=2, cols=2)
            t_ev.rows[0].cells[0].text = "Evidencia de aprendizaje"
            t_ev.rows[0].cells[1].text = _join_or_dash([str(x) for x in (ada.get("evidencias_aprendizaje") or [])])
            t_ev.rows[1].cells[0].text = "Instrumento de evaluacion"
            t_ev.rows[1].cells[1].text = str(ada.get("instrumento_evaluacion", "Lista de cotejo"))

            sesiones = ada.get("sesiones_desarrolladas") or []
            for ses in sesiones:
                ses_num = _num(ses.get("sesion"), "?")
                document.add_paragraph(f"SESION {ses_num}")
                document.add_paragraph(f"Tema: {str(ses.get('tema', '')).strip()}")
                obj_ses = str(ses.get("objetivo_especifico") or ses.get("objetivo") or "").strip()
                document.add_paragraph(f"Objetivo especifico: {obj_ses}")

                resumen = ses.get("resumen_datos_esenciales") or []
                document.add_paragraph("Resumen de los datos esenciales de los temas a tratar:")
                if resumen:
                    _add_list(document, [str(x) for x in resumen])
                else:
                    document.add_paragraph("-")

                for phase in ("inicio", "desarrollo", "cierre"):
                    document.add_paragraph(phase.upper())
                    phase_items = _phase_items(ses, phase)
                    if phase_items:
                        _add_list(document, phase_items)
                    else:
                        document.add_paragraph("Sin actividades registradas.", style="List Bullet")

    document.add_heading("6. Evaluacion de Producto", level=1)
    t_prod = document.add_table(rows=4, cols=5)
    t_prod.rows[0].cells[0].text = "Actividad integradora"
    t_prod.rows[0].cells[1].text = "Duracion en sesiones"
    t_prod.rows[0].cells[2].text = "0"
    t_prod.rows[0].cells[3].text = "Duracion en horas"
    t_prod.rows[0].cells[4].text = "0"
    t_prod.rows[1].cells[0].text = "Actividad integradora"
    t_prod.rows[1].cells[1].text = "Puntaje"
    t_prod.rows[1].cells[2].text = "0"
    t_prod.rows[2].cells[0].text = "Competencia de la asignatura"
    t_prod.rows[2].cells[1].text = str(competencia)
    t_prod.rows[3].cells[0].text = "Estrategias de ensenanza y aprendizaje"
    t_prod.rows[3].cells[1].text = "No aplica en la configuracion actual del curso"

    t_prod_ev = document.add_table(rows=2, cols=2)
    t_prod_ev.rows[0].cells[0].text = "Evidencia de aprendizaje"
    t_prod_ev.rows[0].cells[1].text = "No aplica"
    t_prod_ev.rows[1].cells[0].text = "Instrumento de evaluacion"
    t_prod_ev.rows[1].cells[1].text = "Lista de cotejo"

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))
    return output_docx


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Exporta planeacion didactica JSON a DOCX")
    parser.add_argument("--input-json", required=True, help="Ruta al JSON de salida del pipeline")
    parser.add_argument("--output-docx", help="Ruta destino del archivo DOCX")
    parser.add_argument(
        "--template-docx",
        help="Ruta de plantilla institucional DOCX para rellenar y conservar el formato base",
    )
    parser.add_argument(
        "--start-date",
        help="Fecha de inicio del curso (YYYY-MM-DD, DD/MM/YYYY o '17 de agosto de 2026')",
    )
    parser.add_argument(
        "--end-date",
        help="Fecha de cierre del curso (YYYY-MM-DD, DD/MM/YYYY o '20 de noviembre de 2026')",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_json = Path(args.input_json).resolve()
    if not input_json.exists():
        raise SystemExit(f"No existe el archivo: {input_json}")

    if args.output_docx:
        output_docx = Path(args.output_docx).resolve()
    else:
        stem = input_json.stem
        output_docx = input_json.parent / f"{stem}_planeacion_didactica.docx"

    if args.template_docx:
        template_docx = Path(args.template_docx).resolve()
        if not template_docx.exists():
            raise SystemExit(f"No existe la plantilla DOCX: {template_docx}")
        out = build_doc_from_template(
            input_json,
            template_docx,
            output_docx,
            start_date_text=args.start_date,
            end_date_text=args.end_date,
        )
    else:
        out = build_doc(input_json, output_docx)
    print(str(out))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

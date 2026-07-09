from __future__ import annotations

import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.table import Table
from pypdf import PdfReader

from .program_metadata import extract_program_metadata
from .schemas import ParsedCourseInput


class DocxParsingError(Exception):
    pass


def _normalize(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    return normalized.strip().lower()


def _clean_item(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^[\-\*\u2022\d\.)\s]+", "", text)
    return text.strip()


def _matches_any_heading(normalized_text: str, candidates: tuple[str, ...]) -> bool:
    return any(normalized_text.startswith(candidate) for candidate in candidates)


def _normalize_heading_candidate(normalized_text: str) -> str:
    return re.sub(r"^[\d\.)\s]+", "", normalized_text).strip()


def _extract_sections(lines: list[str]) -> dict[str, list[str]]:
    objective_headings = (
        "objetivos",
        "objetivo",
        "experiencias de aprendizaje",
        "resultados de aprendizaje",
        "proposito",
    )
    competency_headings = (
        "competencias",
        "competencia",
        "competencias del perfil de egreso",
        "perfil de egreso",
    )
    course_competency_heading = "competencia de la asignatura"
    syllabus_headings = (
        "syllabus",
        "temario",
        "contenidos esenciales",
        "contenidos",
        "unidades",
        "temas",
    )
    syllabus_end_headings = (
        "estrategias de ensenanza",
        "estrategias generales de evaluacion",
        "descripcion general de las practicas",
        "practicas formativas",
    )

    sections: dict[str, list[str]] = {
        "objetivos": [],
        "competencias": [],
        "syllabus": [],
        "course_competency": [],
    }
    current_section: str | None = None

    for line in lines:
        raw_text = line.strip()
        if not raw_text:
            continue

        normalized = _normalize(raw_text)
        heading_candidate = _normalize_heading_candidate(normalized)

        if _matches_any_heading(heading_candidate, objective_headings):
            current_section = "objetivos"
            continue
        if heading_candidate.startswith(course_competency_heading):
            current_section = "course_competency"
            continue
        if _matches_any_heading(heading_candidate, competency_headings):
            current_section = "competencias"
            continue
        if _matches_any_heading(heading_candidate, syllabus_headings):
            current_section = "syllabus"
            continue

        if current_section == "syllabus" and _matches_any_heading(heading_candidate, syllabus_end_headings):
            current_section = None
            continue

        if current_section == "course_competency" and (
            _matches_any_heading(heading_candidate, syllabus_headings)
            or _matches_any_heading(heading_candidate, objective_headings)
            or _matches_any_heading(heading_candidate, competency_headings)
        ):
            current_section = None

        if current_section is None:
            continue

        cleaned = _clean_item(raw_text)
        if cleaned:
            sections[current_section].append(cleaned)
            if current_section == "course_competency":
                sections["competencias"].append(cleaned)

    return sections


def _parse_from_docx(file_path: Path) -> dict[str, list[str]]:
    try:
        document = Document(str(file_path))
    except Exception as exc:  # pragma: no cover
        raise DocxParsingError(f"No se pudo abrir el DOCX: {exc}") from exc

    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.extend(part.strip() for part in text.splitlines() if part.strip())

    sections = _extract_sections(lines)
    sections["_planning_context"] = [_extract_docx_planning_context(document.tables)]
    return sections


def _extract_docx_planning_context(tables: list[Table]) -> dict:
    normalized_cells: list[str] = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    normalized_cells.append(_normalize(text))

    period_unit_headers = [
        value for value in normalized_cells if value.startswith("secuencia didactica unidad")
    ]
    activity_headers = [
        value for value in normalized_cells if value.startswith("actividad no")
    ]

    return {
        "detected_units": len(set(period_unit_headers)),
        "detected_activity_blocks": len(activity_headers),
        "has_product_evaluation_section": any(
            value.startswith("evaluacion de producto") for value in normalized_cells
        ),
        "has_integrator_activity": any(
            value.startswith("actividad integradora") for value in normalized_cells
        ),
    }


def _parse_from_pdf(file_path: Path) -> dict[str, list[str]]:
    try:
        reader = PdfReader(str(file_path))
    except Exception as exc:  # pragma: no cover
        raise DocxParsingError(f"No se pudo abrir el PDF: {exc}") from exc

    lines: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        lines.extend(page_text.splitlines())

    sections = _extract_sections(lines)
    sections["_program_metadata"] = [extract_program_metadata("\n".join(lines))]
    return sections


def parse_course_file(file_path: Path, course_name: str) -> ParsedCourseInput:
    if not file_path.exists():
        raise DocxParsingError("El archivo de entrada no existe")

    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        sections = _parse_from_docx(file_path)
    elif suffix == ".pdf":
        sections = _parse_from_pdf(file_path)
    else:
        raise DocxParsingError("Formato no soportado. Usa .docx o .pdf")

    if not sections["objetivos"]:
        raise DocxParsingError("No se encontraron objetivos en el archivo")
    if not sections["competencias"]:
        raise DocxParsingError("No se encontraron competencias en el archivo")
    if not sections["syllabus"]:
        raise DocxParsingError("No se encontro el syllabus/temario en el archivo")

    syllabus_text = "\n".join(sections["syllabus"])
    course_competency_parts = [item.strip() for item in sections.get("course_competency", []) if item.strip()]
    course_competency_text = " ".join(course_competency_parts).strip() if course_competency_parts else None

    program_metadata = sections.get("_program_metadata", [None])[0]
    unidades = None
    if isinstance(program_metadata, dict):
        unidades = program_metadata.get("contenidos_esenciales") or None

    return ParsedCourseInput(
        course_name=course_name,
        objectives=sections["objetivos"],
        competencies=sections["competencias"],
        syllabus=syllabus_text,
        course_competency=course_competency_text,
        planning_context=(sections.get("_planning_context", [None])[0]),
        program_metadata=program_metadata,
        unidades=unidades,
    )


def parse_course_docx(docx_path: Path, course_name: str) -> ParsedCourseInput:
    return parse_course_file(docx_path, course_name)
